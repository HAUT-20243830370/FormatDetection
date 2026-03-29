from docx import Document
import re
import sys
import os
import glob


def check_figure_numbering(doc):
    print('\n' + '=' * 80)
    print('【图序编号检测】')
    print('=' * 80)
    
    current_chapter = None
    chapter_figures = {}
    errors = []
    
    chapter_pattern = re.compile(r'^(\d+)(\.\d+)*\s+')
    figure_pattern1 = re.compile(r'^图(\d+)-(\d+)')
    figure_pattern2 = re.compile(r'^图(\d+)\.(\d+)')
    figure_pattern4 = re.compile(r'^图\s*(\d+)-(\d+)')
    figure_pattern6 = re.compile(r'^图\s*(\d+)\.(\d+)')
    figure_pattern3 = re.compile(r'^图(\d+)\s*')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        chapter_match = chapter_pattern.match(text)
        if chapter_match:
            main_chapter = chapter_match.group(1)
            if main_chapter != current_chapter:
                current_chapter = main_chapter
                if current_chapter not in chapter_figures:
                    chapter_figures[current_chapter] = []
        
        fig_chapter = None
        fig_number = None
        figure_text = None
        
        figure_match = figure_pattern6.match(text)
        if figure_match:
            fig_chapter = figure_match.group(1)
            fig_number = int(figure_match.group(2))
            figure_text = f'图 {fig_chapter}.{fig_number}'
        
        if not figure_match:
            figure_match = figure_pattern2.match(text)
            if figure_match:
                fig_chapter = figure_match.group(1)
                fig_number = int(figure_match.group(2))
                figure_text = f'图{fig_chapter}.{fig_number}'
        
        if not figure_match:
            figure_match = figure_pattern1.match(text)
            if figure_match:
                fig_chapter = figure_match.group(1)
                fig_number = int(figure_match.group(2))
                figure_text = f'图{fig_chapter}-{fig_number}'
        
        if not figure_match:
            figure_match = figure_pattern4.match(text)
            if figure_match:
                fig_chapter = figure_match.group(1)
                fig_number = int(figure_match.group(2))
                figure_text = f'图 {fig_chapter}-{fig_number}'
        
        if not figure_match:
            figure_match = figure_pattern3.match(text)
            if figure_match:
                fig_chapter = current_chapter if current_chapter else '0'
                fig_number = int(figure_match.group(1))
                figure_text = f'图{fig_number}'
        
        if figure_match and fig_chapter and fig_number is not None:
            print(f'  找到图序: {figure_text} (段落 {i+1})')
            
            if current_chapter and fig_chapter != current_chapter and fig_chapter != '0':
                errors.append({
                    'type': '章节不匹配',
                    'paragraph': i + 1,
                    'text': text,
                    'message': f'图序"{figure_text}"所属章节({fig_chapter})与当前章节({current_chapter})不匹配'
                })
            
            if fig_chapter not in chapter_figures:
                chapter_figures[fig_chapter] = []
            chapter_figures[fig_chapter].append({
                'number': fig_number,
                'paragraph': i + 1,
                'text': text,
                'display_text': figure_text
            })
    
    all_valid = True
    for chapter in sorted(chapter_figures.keys(), key=lambda x: int(x)):
        figures = chapter_figures[chapter]
        print(f'\n章节 {chapter} 的图序:')
        
        if not figures:
            print(f'  (无图)')
            continue
        
        if figures[0]['number'] != 1:
            errors.append({
                'type': '起始错误',
                'chapter': chapter,
                'paragraph': figures[0]['paragraph'],
                'text': figures[0]['text'],
                'message': f'章节{chapter}的图序应该从"图{chapter}-1"开始'
            })
            all_valid = False
        
        expected_num = 1
        for fig in figures:
            if fig['number'] != expected_num:
                if fig['number'] < expected_num:
                    errors.append({
                        'type': '重复/倒退',
                        'chapter': chapter,
                        'paragraph': fig['paragraph'],
                        'text': fig['text'],
                        'message': f'章节{chapter}中，期望"图{chapter}-{expected_num}"，但找到"{fig["text"]}"'
                    })
                else:
                    errors.append({
                        'type': '跳序',
                        'chapter': chapter,
                        'paragraph': fig['paragraph'],
                        'text': fig['text'],
                        'message': f'章节{chapter}中，期望"图{chapter}-{expected_num}"，但找到"{fig["text"]}"'
                    })
                all_valid = False
            
            print(f'  {fig["display_text"]} (段落 {fig["paragraph"]})', end='')
            if fig['number'] != expected_num:
                print(' ❌', end='')
            print()
            
            expected_num = fig['number'] + 1
    
    if errors:
        print('\n发现的问题:')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. [{error["type"]}]')
            if 'paragraph' in error:
                print(f'   段落: {error["paragraph"]}')
            if 'text' in error:
                print(f'   内容: {error["text"][:80]}...')
            print(f'   说明: {error["message"]}')
    else:
        print('\n✅ 所有图序编号检查通过！')
    
    return len(errors) == 0


def check_references(doc):
    print('\n' + '=' * 80)
    print('【GB/T 7714—2015参考文献格式检测】')
    print('=' * 80)
    
    in_references = False
    references = []
    citations = set()
    citation_order = []
    citation_locations = {}
    errors = []
    warnings = []
    
    citation_pattern = re.compile(r'\[(\d+(?:[,-]\s*\d+)*)\]')
    ref_number_pattern = re.compile(r'^\[(\d+)\]\s*(.*)')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if ('参考文献' in text or '参 考 文 献' in text) and len(text) < 15:
            in_references = True
            print(f'\n找到参考文献标题 (段落 {i+1}): {text}')
            continue
        
        if not in_references:
            matches = citation_pattern.findall(text)
            for match in matches:
                parts = re.split(r'[,-]', match)
                for part in parts:
                    part = part.strip()
                    if part:
                        ref_num = int(part)
                        citations.add(ref_num)
                        citation_order.append(ref_num)
                        if ref_num not in citation_locations:
                            citation_locations[ref_num] = []
                        citation_locations[ref_num].append({
                            'paragraph': i + 1,
                            'text': text[:60] + '...' if len(text) > 60 else text
                        })
        
        if in_references and text:
            ref_match = ref_number_pattern.match(text)
            if ref_match:
                ref_num = int(ref_match.group(1))
                ref_content = ref_match.group(2).strip()
                references.append({
                    'number': ref_num,
                    'paragraph': i + 1,
                    'content': ref_content,
                    'full_text': text
                })
            else:
                has_ref_marker = False
                for marker in ['[J]', '[D]', '[M]', '[N]', '[P]', '[S]', '[EB/OL]']:
                    if marker in text:
                        has_ref_marker = True
                        break
                
                if has_ref_marker:
                    ref_num = len(references) + 1
                    references.append({
                        'number': ref_num,
                        'paragraph': i + 1,
                        'content': text,
                        'full_text': text
                    })
                else:
                    continue
            
            if len(references) > 0:
                ref_info = references[-1]
                ref_num = ref_info['number']
                ref_content = ref_info['content']
                
                ref_type = None
                
                if '[J]' in ref_content:
                    ref_type = '期刊文章'
                    print(f'[{ref_num}] {ref_type}: {ref_content[:60]}...')
                    
                    if not re.search(r'\[J\]\.', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '期刊文章格式：[J]后应直接跟英文句号，应为 [J].'
                        })
                    
                    if re.search(r'\[J\]。', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '期刊文章格式：不应使用中文句号，应为英文句号 .'
                        })
                    
                    if re.search(r'，', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '期刊文章格式：不应使用中文逗号，应为英文逗号 ,'
                        })
                
                elif '[D]' in ref_content:
                    ref_type = '学位论文'
                    print(f'[{ref_num}] {ref_type}: {ref_content[:60]}...')
                    
                    if not re.search(r'\[D\]\.', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '学位论文格式：[D]后应直接跟英文句号，应为 [D].'
                        })
                    
                    if re.search(r'\[D\]。', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '学位论文格式：不应使用中文句号，应为英文句号 .'
                        })
                    
                    if re.search(r'，', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '学位论文格式：不应使用中文逗号，应为英文逗号 ,'
                        })
                
                elif '[M]' in ref_content:
                    ref_type = '图书'
                    print(f'[{ref_num}] {ref_type}: {ref_content[:60]}...')
                    
                    if not re.search(r'\[M\]\.', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '图书格式：[M]后应直接跟英文句号，应为 [M].'
                        })
                    
                    if re.search(r'\[M\]。', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '图书格式：不应使用中文句号，应为英文句号 .'
                        })
                    
                    if re.search(r'，', ref_content):
                        errors.append({
                            'type': '格式错误',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '图书格式：不应使用中文逗号，应为英文逗号 ,'
                        })
                
                else:
                    ref_type = '其他类型'
                    print(f'[{ref_num}] {ref_type}: {ref_content[:60]}...')
                    if not re.search(r'\[.*?\]', ref_content):
                        errors.append({
                            'type': '缺少文献类型标识',
                            'number': ref_num,
                            'paragraph': i + 1,
                            'content': text,
                            'message': '缺少文献类型标识([J]/[D]/[M]等)'
                        })
    
    if references:
        ref_numbers = [ref['number'] for ref in references]
        ref_set = set(ref_numbers)
        cited_but_not_listed = sorted(citations - ref_set)
        listed_but_not_cited = sorted(ref_set - citations)
        
        print(f'\n参考文献总数: {len(references)}')
        print(f'正文引用标记数: {len(citations)}')
        
        if cited_but_not_listed:
            errors.append({
                'type': '引用缺失',
                'message': f'正文引用了但参考文献列表中缺少: {cited_but_not_listed}'
            })
        
        if listed_but_not_cited:
            print(f'\n发现 {len(listed_but_not_cited)} 个未被引用的参考文献:')
            for ref_num in listed_but_not_cited:
                ref_info = next((ref for ref in references if ref['number'] == ref_num), None)
                if ref_info:
                    print(f'\n  [{ref_num}] 段落 {ref_info["paragraph"]}:')
                    print(f'    内容: {ref_info["full_text"]}')
                    errors.append({
                        'type': '未被引用',
                        'number': ref_num,
                        'paragraph': ref_info['paragraph'],
                        'content': ref_info['full_text'],
                        'message': f'参考文献 [{ref_num}] 在正文中未被引用'
                    })
        
        if citation_order:
            print(f'\n正文引用顺序: {citation_order}')
            
            last_cited = 0
            for idx, ref_num in enumerate(citation_order):
                if ref_num <= last_cited:
                    errors.append({
                        'type': '引用顺序错误',
                        'message': f'第{idx + 1}次引用 [{ref_num}] 出现在 [{last_cited}] 之后，未按从小到大顺序引用'
                    })
                last_cited = ref_num
        
        for ref_num, locations in citation_locations.items():
            if len(locations) > 1:
                location_str = ', '.join([f'段落{loc["paragraph"]}' for loc in locations])
                errors.append({
                    'type': '重复引用',
                    'message': f'参考文献 [{ref_num}] 被重复引用了 {len(locations)} 次 ({location_str})'
                })
    
    if errors:
        print('\n发现的问题:')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. [{error["type"]}]')
            if 'paragraph' in error:
                print(f'   段落: {error["paragraph"]}')
            if 'content' in error:
                print(f'   内容: {error["content"][:80]}...')
            print(f'   说明: {error["message"]}')
    else:
        print('\n✅ GB/T 7714—2015参考文献格式检查通过！')
    
    if warnings:
        print('\n⚠️  建议:')
        for i, warning in enumerate(warnings, 1):
            print(f'\n{i}. [{warning["type"]}]')
            if 'paragraph' in warning:
                print(f'   段落: {warning["paragraph"]}')
            if 'content' in warning:
                print(f'   内容: {warning["content"][:80]}...')
            print(f'   说明: {warning["message"]}')
    
    return len(errors) == 0


def check_reference_crossref(doc):
    print('\n' + '=' * 80)
    print('【参考文献交叉引用检测】')
    print('=' * 80)
    
    import re
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    in_references = False
    references = []
    citations = set()
    citation_details = {}
    
    citation_pattern = re.compile(r'\[(\d+)\]')
    ref_number_pattern = re.compile(r'^\[(\d+)\]\s*')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not in_references:
            if ('参考文献' in text or '参 考 文 献' in text) and len(text) < 20:
                in_references = True
                print(f'\n找到参考文献标题 (段落 {i+1}): {text}')
                continue
            
            matches = citation_pattern.findall(text)
            for match in matches:
                ref_num = int(match)
                citations.add(ref_num)
                
                if ref_num not in citation_details:
                    citation_details[ref_num] = []
                
                has_hyperlink = False
                has_instrtext = False
                
                para_elem = para._element
                hyperlinks = para_elem.findall('.//w:hyperlink', namespaces=ns)
                if hyperlinks:
                    has_hyperlink = True
                
                instrtexts = para_elem.findall('.//w:instrText', namespaces=ns)
                for instr in instrtexts:
                    if instr.text and ('REF' in instr.text or 'PAGEREF' in instr.text):
                        has_instrtext = True
                        break
                
                citation_details[ref_num].append({
                    'paragraph': i + 1,
                    'text': text[:60] + '...' if len(text) > 60 else text,
                    'has_hyperlink': has_hyperlink,
                    'has_instrtext': has_instrtext
                })
        else:
            if (text.startswith('致谢') or text.startswith('致 谢') or 
                text.startswith('结 论') or text.startswith('结论')):
                print(f'\n遇到结束标记，停止检测 (段落 {i+1}): {text}')
                break
            
            if not text:
                continue
            
            ref_match = ref_number_pattern.match(text)
            if ref_match:
                ref_num = int(ref_match.group(1))
                references.append(ref_num)
            else:
                has_ref_marker = False
                for marker in ['[J]', '[D]', '[M]', '[N]', '[P]', '[S]', '[EB/OL]']:
                    if marker in text:
                        has_ref_marker = True
                        break
                if has_ref_marker:
                    ref_num = len(references) + 1
                    references.append(ref_num)
    
    print(f'\n正文中引用的参考文献: {sorted(citations)}')
    print(f'参考文献列表中的文献: {sorted(references)}')
    
    errors = []
    all_cited_ok = True
    
    for ref_num in sorted(citations):
        if ref_num not in citation_details:
            continue
        
        details = citation_details[ref_num]
        all_crossref_ok = True
        
        for j, detail in enumerate(details):
            if not detail['has_hyperlink'] and not detail['has_instrtext']:
                all_crossref_ok = False
                all_cited_ok = False
                errors.append({
                    'type': '缺少交叉引用',
                    'number': ref_num,
                    'paragraph': detail['paragraph'],
                    'content': detail['text'],
                    'message': f'参考文献 [{ref_num}] 的第 {j+1} 次引用未设置交叉引用'
                })
        
        if all_crossref_ok:
            print(f'\n[{ref_num}] ✅ 所有引用均已设置交叉引用')
        else:
            print(f'\n[{ref_num}] ❌ 部分引用未设置交叉引用')
    
    if errors:
        print('\n发现的问题:')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. [{error["type"]}]')
            print(f'   参考文献: [{error["number"]}]')
            print(f'   段落: {error["paragraph"]}')
            print(f'   内容: {error["content"]}')
            print(f'   说明: {error["message"]}')
    else:
        print('\n✅ 所有参考文献引用均已设置交叉引用！')
    
    return all_cited_ok


def check_indent(doc):
    print('\n' + '=' * 80)
    print('【段落首行缩进检测】')
    print('=' * 80)
    
    errors = []
    total_chinese_paras = 0
    total_correct = 0
    
    chinese_pattern = re.compile(r'[\u4e00-\u9fff]')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        if not chinese_pattern.search(text):
            continue
        
        if i < 50:
            if (text.startswith('教育培训学院') or 
                text.startswith('2026届') or
                text.startswith('题    目') or
                text.startswith('学习形式') or
                text.startswith('学习层次') or
                text.startswith('专    业') or
                text.startswith('学生姓名') or
                text.startswith('学    号') or
                text.startswith('指导教师') or
                text.startswith('答辩日期') or
                text.startswith('评价数据的')):
                continue
        
        if (text.startswith('第') and '章' in text) or \
           (text.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9')) and ' ' in text[:10]) or \
           (text.startswith('摘要') or text.startswith('Abstract') or 
            text.startswith('关键词') or text.startswith('Keywords') or
            text.startswith('参考文献') or text.startswith('参 考 文 献') or
            text.startswith('致 谢') or text.startswith('致谢') or
            text.startswith('目 次') or text.startswith('目次') or
            text.startswith('结 论') or text.startswith('结论') or
            text.startswith('摘  要') or text.startswith('目   次')):
            continue
        
        if text.startswith('图 ') or text.startswith('表 ') or re.match(r'^\d+\.\d+', text):
            continue
        
        if re.match(r'^\[\d+\]', text):
            continue
        
        total_chinese_paras += 1
        
        indent_value = 0.0
        para_elem = para._element
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        pPr = para_elem.find('./w:pPr', namespaces=ns)
        
        if pPr is not None:
            ind = pPr.find('./w:ind', namespaces=ns)
            if ind is not None:
                first_line_chars = ind.get(f'{{{ns["w"]}}}firstLineChars')
                if first_line_chars is not None:
                    indent_value = int(first_line_chars) / 100.0
                else:
                    first_line = ind.get(f'{{{ns["w"]}}}firstLine')
                    if first_line is not None:
                        indent_value = int(first_line) / 152400
        
        expected_indent = 2.0
        tolerance = 0.5
        
        if abs(indent_value - expected_indent) > tolerance:
            errors.append({
                'paragraph': i + 1,
                'text': text[:50] + '...' if len(text) > 50 else text,
                'actual_indent': round(indent_value, 2),
                'expected_indent': expected_indent
            })
        else:
            total_correct += 1
    
    print(f'\n中文段落总数: {total_chinese_paras}')
    print(f'首行缩进正确: {total_correct}')
    print(f'首行缩进错误: {len(errors)}')
    
    if errors:
        print(f'\n发现 {len(errors)} 个首行缩进问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 段落 {error["paragraph"]}:')
            print(f'   内容: {error["text"]}')
            print(f'   实际缩进: {error["actual_indent"]} 字符')
            print(f'   期望缩进: {error["expected_indent"]} 字符')
    else:
        print(f'\n✅ 所有中文段落首行缩进检查通过！')
        print(f'   所有段落均设置为 {expected_indent} 字符首行缩进')
    
    return len(errors) == 0


def check_chinese_font(doc):
    print('\n' + '=' * 80)
    print('【中文字体检测 - 详细信息】')
    print('=' * 80)
    
    errors = []
    total_chinese_paras = 0
    total_correct = 0
    
    chinese_pattern = re.compile(r'[\u4e00-\u9fff]')
    
    def is_font_songti(font_name):
        if not font_name:
            return False
        return ('宋体' in font_name or 
                'Songti' in font_name or 
                'SimSun' in font_name)
    
    def get_eastasia_font_from_style(style, ns):
        if style is None:
            return None
        style_elem = style._element
        rPr = style_elem.find('./w:rPr', namespaces=ns)
        if rPr is not None:
            rFonts = rPr.find('./w:rFonts', namespaces=ns)
            if rFonts is not None:
                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                if eastAsia:
                    return eastAsia
        based_on = getattr(style, 'based_on', None)
        if based_on:
            return get_eastasia_font_from_style(based_on, ns)
        return None
    
    def get_font_from_element(element, ns):
        if element is None:
            return None
        rFonts = element.find('./w:rPr/w:rFonts', namespaces=ns)
        if rFonts is not None:
            eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
            if eastAsia:
                return eastAsia
            ascii = rFonts.get(f'{{{ns["w"]}}}ascii')
            if ascii:
                return ascii
            hAnsi = rFonts.get(f'{{{ns["w"]}}}hAnsi')
            if hAnsi:
                return hAnsi
        return None
    
    def get_font_from_style(style, ns):
        if style is None:
            return None
        style_elem = style._element
        font = get_font_from_element(style_elem, ns)
        if font:
            return font
        based_on = getattr(style, 'based_on', None)
        if based_on:
            return get_font_from_style(based_on, ns)
        return None
    
    def has_eastasia_hint(run_element, ns):
        rPr = run_element.find('./w:rPr', namespaces=ns)
        if rPr is not None:
            rFonts = rPr.find('./w:rFonts', namespaces=ns)
            if rFonts is not None:
                hint = rFonts.get(f'{{{ns["w"]}}}hint')
                return hint == 'eastAsia'
        return False
    
    def get_run_font(run, para, doc_styles, ns):
        run_elem = run._element
        
        rPr = run_elem.find('./w:rPr', namespaces=ns)
        if rPr is not None:
            rFonts = rPr.find('./w:rFonts', namespaces=ns)
            if rFonts is not None:
                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                if eastAsia:
                    return eastAsia
        
        if has_eastasia_hint(run_elem, ns):
            if para.style:
                eastAsia_font = get_eastasia_font_from_style(para.style, ns)
                if eastAsia_font:
                    return eastAsia_font
            default_style = doc_styles.get('Normal')
            if default_style:
                eastAsia_font = get_eastasia_font_from_style(default_style, ns)
                if eastAsia_font:
                    return eastAsia_font
        
        if para.style:
            eastAsia_font = get_eastasia_font_from_style(para.style, ns)
            if eastAsia_font:
                return eastAsia_font
        
        default_style = doc_styles.get('Normal')
        if default_style:
            eastAsia_font = get_eastasia_font_from_style(default_style, ns)
            if eastAsia_font:
                return eastAsia_font
        
        font = get_font_from_element(run_elem, ns)
        if font:
            return font
        
        if para.style:
            font = get_font_from_style(para.style, ns)
            if font:
                return font
        
        if default_style:
            font = get_font_from_style(default_style, ns)
            if font:
                return font
        
        return None
    
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    doc_styles = {s.name: s for s in doc.styles}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        if not chinese_pattern.search(text):
            continue
        
        if i < 50:
            if (text.startswith('教育培训学院') or 
                text.startswith('2026届') or
                text.startswith('题    目') or
                text.startswith('学习形式') or
                text.startswith('学习层次') or
                text.startswith('专    业') or
                text.startswith('学生姓名') or
                text.startswith('学    号') or
                text.startswith('指导教师') or
                text.startswith('答辩日期') or
                text.startswith('评价数据的')):
                continue
        
        if (text.startswith('第') and '章' in text) or \
           (text.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9')) and ' ' in text[:10]) or \
           (text.startswith('摘要') or text.startswith('Abstract') or 
            text.startswith('关键词') or text.startswith('Keywords') or
            text.startswith('参考文献') or text.startswith('参 考 文 献') or
            text.startswith('致 谢') or text.startswith('致谢') or
            text.startswith('目 次') or text.startswith('目次') or
            text.startswith('结 论') or text.startswith('结论') or
            text.startswith('摘  要') or text.startswith('目   次')):
            continue
        
        if text.startswith('图 ') or text.startswith('表 ') or re.match(r'^\d+\.\d+', text):
            continue
        
        if re.match(r'^\[\d+\]', text):
            continue
        
        total_chinese_paras += 1
        
        has_chinese_error = False
        found_chinese_text = False
        char_font_info = []
        
        for run_idx, run in enumerate(para.runs):
            run_text = run.text
            
            if not run_text:
                continue
            
            if run.font.superscript or run.font.subscript:
                continue
            
            current_font = get_run_font(run, para, doc_styles, ns)
            
            for char in run_text:
                if chinese_pattern.search(char):
                    found_chinese_text = True
                    char_font_info.append({
                        'char': char,
                        'font': current_font or '未知',
                        'is_songti': is_font_songti(current_font)
                    })
                    if not is_font_songti(current_font):
                        has_chinese_error = True
        
        if found_chinese_text and has_chinese_error:
            errors.append({
                'paragraph': i + 1,
                'text': text[:80] + '...' if len(text) > 80 else text,
                'char_font_info': char_font_info
            })
        else:
            total_correct += 1
    
    print(f'\n中文段落总数: {total_chinese_paras}')
    print(f'字体正确: {total_correct}')
    print(f'字体错误: {len(errors)}')
    print(f'（表格字体检测已跳过）')
    
    if errors:
        print(f'\n发现 {len(errors)} 个段落字体问题:\n')
        for i, error in enumerate(errors[:10], 1):
            print(f'\n{i}. 段落 {error["paragraph"]}:')
            print(f'   内容: {error["text"]}')
            print(f'   逐字字体信息 (前30字):')
            char_info_str = ''
            for j, char_info in enumerate(error['char_font_info'][:30], 1):
                if not char_info['is_songti']:
                    char_info_str += f"『{char_info['char']}』({char_info['font']}) "
                else:
                    char_info_str += f"{char_info['char']} "
            print(f'   {char_info_str}')
            if len(error['char_font_info']) > 30:
                print(f'   ... 还有 {len(error["char_font_info"]) - 30} 个字符')
        if len(errors) > 10:
            print(f'\n... 还有 {len(errors) - 10} 个段落问题未显示')
    
    if not errors:
        print(f'\n✅ 所有中文段落字体检查通过！')
    
    return len(errors) == 0


def check_toc_heading1_font(doc):
    print('\n' + '=' * 80)
    print('【目录字体检测】')
    print('=' * 80)
    
    errors = []
    found_headings = []
    
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    doc_styles = {s.name: s for s in doc.styles}
    
    def is_font_songti(font_name):
        if not font_name:
            return False
        return ('宋体' in font_name or 
                'Songti' in font_name or 
                'SimSun' in font_name)
    
    def is_font_times_new_roman(font_name):
        if not font_name:
            return False
        return ('Times New Roman' in font_name or 
                'TimesNewRoman' in font_name)
    
    def get_eastasia_font_from_style(style, ns):
        if style is None:
            return None
        style_elem = style._element
        rPr = style_elem.find('./w:rPr', namespaces=ns)
        if rPr is not None:
            rFonts = rPr.find('./w:rFonts', namespaces=ns)
            if rFonts is not None:
                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                if eastAsia:
                    return eastAsia
        based_on = getattr(style, 'based_on', None)
        if based_on:
            return get_eastasia_font_from_style(based_on, ns)
        return None
    
    def get_ascii_font_from_style(style, ns):
        if style is None:
            return None
        style_elem = style._element
        rPr = style_elem.find('./w:rPr', namespaces=ns)
        if rPr is not None:
            rFonts = rPr.find('./w:rFonts', namespaces=ns)
            if rFonts is not None:
                ascii_font = rFonts.get(f'{{{ns["w"]}}}ascii')
                if ascii_font:
                    return ascii_font
        based_on = getattr(style, 'based_on', None)
        if based_on:
            return get_ascii_font_from_style(based_on, ns)
        return None
    
    root = doc.element.body
    for i, p_elem in enumerate(root.findall('.//w:p', namespaces=ns)):
        text_parts = []
        for t in p_elem.findall('.//w:t', namespaces=ns):
            if t.text:
                text_parts.append(t.text)
        text = ''.join(text_parts)
        
        if not text:
            continue
        
        style_elem = p_elem.find('.//w:pStyle', namespaces=ns)
        style_val = style_elem.get(f'{{{ns["w"]}}}val') if style_elem is not None else None
        
        if style_val in ['12', '13']:
            heading_level = '一级' if style_val == '12' else '二级'
            found_headings.append({
                'paragraph': i + 1,
                'text': text,
                'style_id': style_val,
                'level': heading_level
            })
            
            chinese_font = None
            english_font_explicit = None
            
            def has_chinese(text):
                for char in text:
                    if '\u4e00' <= char <= '\u9fff':
                        return True
                return False
            
            def has_english(text):
                for char in text:
                    if char.isalpha():
                        return True
                return False
            
            def has_english_or_digit(text):
                for char in text:
                    if char.isalpha() or char.isdigit():
                        return True
                return False
            
            def has_alpha_in_paragraph(text):
                for char in text:
                    if 'a' <= char <= 'z' or 'A' <= char <= 'Z':
                        return True
                return False
            
            def get_fonts_from_rStyle(rStyle_val, doc_styles, ns):
                for s_name, s in doc_styles.items():
                    s_elem = s._element
                    s_id_elem = s_elem.find('./w:styleId', namespaces=ns)
                    if s_id_elem is not None and s_id_elem.get(f'{{{ns["w"]}}}val') == rStyle_val:
                        rPr = s_elem.find('./w:rPr', namespaces=ns)
                        if rPr is not None:
                            rFonts = rPr.find('./w:rFonts', namespaces=ns)
                            if rFonts is not None:
                                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                                ascii_font = rFonts.get(f'{{{ns["w"]}}}ascii')
                                return eastAsia, ascii_font
                        based_on = s.based_on
                        if based_on:
                            return get_fonts_from_based_on_style(based_on, ns)
                return None, None
            
            def get_fonts_from_based_on_style(style, ns):
                style_elem = style._element
                rPr = style_elem.find('./w:rPr', namespaces=ns)
                if rPr is not None:
                    rFonts = rPr.find('./w:rFonts', namespaces=ns)
                    if rFonts is not None:
                        eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                        ascii_font = rFonts.get(f'{{{ns["w"]}}}ascii')
                        return eastAsia, ascii_font
                based_on = getattr(style, 'based_on', None)
                if based_on:
                    return get_fonts_from_based_on_style(based_on, ns)
                return None, None
            
            r_elems = p_elem.findall('./w:r', namespaces=ns)
            for r_elem in r_elems:
                t_elems = r_elem.findall('./w:t', namespaces=ns)
                run_text = ''.join([t.text for t in t_elems if t.text])
                
                if run_text.strip():
                    rPr = r_elem.find('./w:rPr', namespaces=ns)
                    if rPr is not None:
                        rStyle_elem = rPr.find('./w:rStyle', namespaces=ns)
                        if rStyle_elem is not None:
                            rStyle_val = rStyle_elem.get(f'{{{ns["w"]}}}val')
                            if rStyle_val:
                                eastAsia, ascii_font = get_fonts_from_rStyle(rStyle_val, doc_styles, ns)
                                if has_chinese(run_text) and not chinese_font and eastAsia:
                                    chinese_font = eastAsia
                                if has_english_or_digit(run_text) and not english_font_explicit and ascii_font:
                                    english_font_explicit = ascii_font
                        
                        rFonts = rPr.find('./w:rFonts', namespaces=ns)
                        if rFonts is not None:
                            if has_chinese(run_text) and not chinese_font:
                                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                                if eastAsia:
                                    chinese_font = eastAsia
                            if has_english_or_digit(run_text) and not english_font_explicit:
                                ascii_font = rFonts.get(f'{{{ns["w"]}}}ascii')
                                if ascii_font:
                                    english_font_explicit = ascii_font
            
            english_font = english_font_explicit
            paragraph_has_alpha = has_alpha_in_paragraph(text)
            
            style_obj = None
            for s in doc.styles:
                s_elem = s._element
                s_id_elem = s_elem.find('./w:styleId', namespaces=ns)
                if s_id_elem is not None and s_id_elem.get(f'{{{ns["w"]}}}val') == style_val:
                    style_obj = s
                    break
            
            if style_obj:
                if not chinese_font:
                    chinese_font = get_eastasia_font_from_style(style_obj, ns)
            
            if not chinese_font:
                default_style = doc_styles.get('Normal')
                if default_style:
                    chinese_font = get_eastasia_font_from_style(default_style, ns)
            
            print(f'\n目录{heading_level}标题 (段落 {i + 1}): {text}')
            print(f'  样式ID: {style_val}')
            print(f'  中文字体: {chinese_font or "未知"}')
            print(f'  英文字体: {english_font or "使用段落默认"}')
            
            if not is_font_songti(chinese_font):
                errors.append({
                    'paragraph': i + 1,
                    'text': text,
                    'style_id': style_val,
                    'level': heading_level,
                    'type': '中文字体错误',
                    'actual': chinese_font or '未知',
                    'expected': '宋体'
                })
            
            if paragraph_has_alpha and english_font and not is_font_times_new_roman(english_font):
                errors.append({
                    'paragraph': i + 1,
                    'text': text,
                    'style_id': style_val,
                    'level': heading_level,
                    'type': '英文字体错误',
                    'actual': english_font or '未知',
                    'expected': 'Times New Roman'
                })
    
    if not found_headings:
        print('\n⚠️  未找到目录标题，请确保文档包含自动生成的目录（使用WPS"引用"→"目录"功能）')
        errors.append({
            'type': '未找到',
            'message': '文档中未找到自动生成的目录标题（styleId: 12或13）'
        })
    else:
        print(f'\n找到 {len(found_headings)} 个目录标题（一级和二级）')
    
    if errors:
        print(f'\n发现 {len(errors)} 个目录字体问题:\n')
        for i, error in enumerate(errors, 1):
            if 'paragraph' in error:
                print(f'\n{i}. 段落 {error["paragraph"]}:')
                print(f'   内容: {error["text"]}')
                print(f'   级别: {error["level"]}')
                print(f'   类型: {error["type"]}')
                print(f'   实际字体: {error["actual"]}')
                print(f'   期望字体: {error["expected"]}')
            else:
                print(f'\n{i}. [{error["type"]}]')
                print(f'   说明: {error["message"]}')
    else:
        print(f'\n✅ 所有目录字体检查通过！')
    
    return len(errors) == 0


def check_paragraph_spacing_after(doc):
    print('\n' + '=' * 80)
    print('【段落段后间距检测】')
    print('=' * 80)
    
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    import re
    
    root = doc.element.body
    for i, p_elem in enumerate(root.findall('.//w:p', namespaces=ns)):
        text_parts = []
        for t in p_elem.findall('.//w:t', namespaces=ns):
            if t.text:
                text_parts.append(t.text)
        text = ''.join(text_parts)
        
        if not text:
            continue
        
        if (text.startswith('教育培训学院') or 
            text.startswith('2026届') or
            text.startswith('题    目') or
            text.startswith('学习形式') or
            text.startswith('学习层次') or
            text.startswith('专    业') or
            text.startswith('学生姓名') or
            text.startswith('学    号') or
            text.startswith('指导教师') or
            text.startswith('答辩日期') or
            text.startswith('评价数据的')):
            continue
        
        if (text.startswith('第') and '章' in text) or \
           (text.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9')) and ' ' in text[:10]) or \
           (text.startswith('摘要') or text.startswith('Abstract') or 
            text.startswith('关键词') or text.startswith('Keywords') or
            text.startswith('参考文献') or text.startswith('参 考 文 献') or
            text.startswith('致 谢') or text.startswith('致谢') or
            text.startswith('目 次') or text.startswith('目次') or
            text.startswith('结 论') or text.startswith('结论') or
            text.startswith('摘  要') or text.startswith('目   次')):
            continue
        
        if text.startswith('图 ') or text.startswith('表 ') or re.match(r'^\d+\.\d+', text):
            continue
        
        if re.match(r'^\[\d+\]', text):
            continue
        
        pPr = p_elem.find('./w:pPr', namespaces=ns)
        if pPr is not None:
            spacing = pPr.find('./w:spacing', namespaces=ns)
            if spacing is not None:
                after = spacing.get(f'{{{ns["w"]}}}after')
                afterLines = spacing.get(f'{{{ns["w"]}}}afterLines')
                
                if after is not None or afterLines is not None:
                    if afterLines is not None:
                        afterLines_val = float(afterLines) / 240.0  # 240 twip = 1 行
                        if abs(afterLines_val) > 0.01:
                            errors.append({
                                'paragraph': i + 1,
                                'text': text[:80],
                                'type': '段后间距错误',
                                'actual_units': '行',
                                'actual_value': afterLines_val,
                                'expected_units': '行',
                                'expected_value': 0.0
                            })
                    if after is not None:
                        after_pt = float(after) / 20.0  # 20 twip = 1 磅
                        if abs(after_pt) > 0.01:
                            errors.append({
                                'paragraph': i + 1,
                                'text': text[:80],
                                'type': '段后间距错误',
                                'actual_units': '磅',
                                'actual_value': after_pt,
                                'expected_units': '磅',
                                'expected_value': 0.0
                            })
    
    if errors:
        print(f'\n发现 {len(errors)} 个段落段后间距问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 段落 {error["paragraph"]}:')
            print(f'   内容: {error["text"]}...')
            print(f'   类型: {error["type"]}')
            print(f'   实际: {error["actual_value"]:.1f} {error["actual_units"]}')
            print(f'   期望: {error["expected_value"]:.1f} {error["expected_units"]}')
    else:
        print(f'\n✅ 所有段落段后间距检查通过！')
    
    return len(errors) == 0


def check_body_font_size(doc):
    print('\n' + '=' * 80)
    print('【正文字体大小和行距检测】')
    print('=' * 80)
    
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    import re
    
    root = doc.element.body
    for i, p_elem in enumerate(root.findall('.//w:p', namespaces=ns)):
        text_parts = []
        for t in p_elem.findall('.//w:t', namespaces=ns):
            if t.text:
                text_parts.append(t.text)
        text = ''.join(text_parts)
        
        if not text:
            continue
        
        if i < 100:
            continue
        
        if (text.startswith('第') and '章' in text) or \
           (text.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9')) and ' ' in text[:10]) or \
           (text.startswith('摘要') or text.startswith('Abstract') or 
            text.startswith('关键词') or text.startswith('Keywords') or
            text.startswith('参考文献') or text.startswith('参 考 文 献') or
            text.startswith('致 谢') or text.startswith('致谢') or
            text.startswith('结 论') or text.startswith('结论') or
            text.startswith('摘  要') or text.startswith('目   次')):
            continue
        
        if text.startswith('图 ') or text.startswith('表 ') or re.match(r'^\d+\.\d+', text):
            continue
        
        if re.match(r'^\[\d+\]', text):
            continue
        
        pPr = p_elem.find('./w:pPr', namespaces=ns)
        if pPr is not None:
            spacing = pPr.find('./w:spacing', namespaces=ns)
            if spacing is not None:
                line = spacing.get(f'{{{ns["w"]}}}line')
                lineRule = spacing.get(f'{{{ns["w"]}}}lineRule')
                if line is not None:
                    line_val = float(line) / 240.0
                    if abs(line_val - 1.5) > 0.01:
                        errors.append({
                            'paragraph': i + 1,
                            'text': text[:80],
                            'type': '行距错误',
                            'actual': f'{line_val:.1f} 倍',
                            'expected': '1.5 倍'
                        })
        
        r_elems = p_elem.findall('./w:r', namespaces=ns)
        for j, r_elem in enumerate(r_elems):
            t_elems = r_elem.findall('./w:t', namespaces=ns)
            run_text = ''.join([t.text for t in t_elems if t.text])
            
            if not run_text.strip():
                continue
            
            rPr = r_elem.find('./w:rPr', namespaces=ns)
            if rPr is not None:
                sz = rPr.find('./w:sz', namespaces=ns)
                rFonts = rPr.find('./w:rFonts', namespaces=ns)
                
                if sz is not None:
                    sz_val = float(sz.get(f'{{{ns["w"]}}}val')) / 2.0
                    if abs(sz_val - 12.0) > 0.1:
                        has_chinese = any('\u4e00' <= char <= '\u9fff' for char in run_text)
                        if has_chinese:
                            errors.append({
                                'paragraph': i + 1,
                                'text': text[:80],
                                'type': '字体大小错误',
                                'actual': f'{sz_val:.1f} 磅',
                                'expected': '12.0 磅 (小4号)'
                            })
                
                if rFonts is not None:
                    eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                    ascii_font = rFonts.get(f'{{{ns["w"]}}}ascii')
                    
                    has_chinese = any('\u4e00' <= char <= '\u9fff' for char in run_text)
                    has_english = any('a' <= char <= 'z' or 'A' <= char <= 'Z' for char in run_text)
                    
                    if has_chinese and eastAsia and eastAsia != '宋体':
                        errors.append({
                            'paragraph': i + 1,
                            'text': text[:80],
                            'type': '中文字体错误',
                            'actual': eastAsia,
                            'expected': '宋体'
                        })
                    
                    if has_english and ascii_font and 'Times New Roman' not in ascii_font:
                        errors.append({
                            'paragraph': i + 1,
                            'text': text[:80],
                            'type': '英文字体错误',
                            'actual': ascii_font,
                            'expected': 'Times New Roman'
                        })
    
    if errors:
        print(f'\n发现 {len(errors)} 个正文字体大小和行距问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 段落 {error["paragraph"]}:')
            print(f'   内容: {error["text"]}...')
            print(f'   类型: {error["type"]}')
            print(f'   实际: {error["actual"]}')
            print(f'   期望: {error["expected"]}')
    else:
        print(f'\n✅ 所有正文字体大小和行距检查通过！')
    
    return len(errors) == 0


def check_table_figure_font(doc):
    print('\n' + '=' * 80)
    print('【表格和插图字体检测】')
    print('=' * 80)
    
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    if not para.text.strip():
                        continue
                    
                    para_elem = para._element
                    r_elems = para_elem.findall('.//w:r', namespaces=ns)
                    for r_elem in r_elems:
                        t_elems = r_elem.findall('./w:t', namespaces=ns)
                        run_text = ''.join([t.text for t in t_elems if t.text])
                        
                        if not run_text.strip():
                            continue
                        
                        rPr = r_elem.find('./w:rPr', namespaces=ns)
                        if rPr is not None:
                            sz = rPr.find('./w:sz', namespaces=ns)
                            rFonts = rPr.find('./w:rFonts', namespaces=ns)
                            
                            if sz is not None:
                                sz_val = float(sz.get(f'{{{ns["w"]}}}val')) / 2.0
                                if abs(sz_val - 10.5) > 0.1:
                                    has_chinese = any('\u4e00' <= char <= '\u9fff' for char in run_text)
                                    if has_chinese:
                                        errors.append({
                                            'table': table_idx + 1,
                                            'cell': f'({row_idx+1},{cell_idx+1})',
                                            'text': para.text[:60],
                                            'type': '字体大小错误',
                                            'actual': f'{sz_val:.1f} 磅',
                                            'expected': '10.5 磅 (5号)'
                                        })
                            
                            if rFonts is not None:
                                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                                ascii_font = rFonts.get(f'{{{ns["w"]}}}ascii')
                                
                                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in run_text)
                                has_english = any('a' <= char <= 'z' or 'A' <= char <= 'Z' for char in run_text)
                                
                                if has_chinese and eastAsia and eastAsia != '宋体':
                                    errors.append({
                                        'table': table_idx + 1,
                                        'cell': f'({row_idx+1},{cell_idx+1})',
                                        'text': para.text[:60],
                                        'type': '中文字体错误',
                                        'actual': eastAsia,
                                        'expected': '宋体'
                                    })
                                
                                if has_english and ascii_font and 'Times New Roman' not in ascii_font:
                                    errors.append({
                                        'table': table_idx + 1,
                                        'cell': f'({row_idx+1},{cell_idx+1})',
                                        'text': para.text[:60],
                                        'type': '英文字体错误',
                                        'actual': ascii_font,
                                        'expected': 'Times New Roman'
                                    })
    
    if errors:
        print(f'\n发现 {len(errors)} 个表格字体问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 表格 {error["table"]} 单元格 {error["cell"]}:')
            print(f'   内容: {error["text"]}...')
            print(f'   类型: {error["type"]}')
            print(f'   实际: {error["actual"]}')
            print(f'   期望: {error["expected"]}')
    else:
        print(f'\n✅ 所有表格和插图字体检查通过！')
    
    return len(errors) == 0


def check_reference_count(doc):
    print('\n' + '=' * 80)
    print('【参考文献数量检测】')
    print('=' * 80)
    
    import re
    
    reference_count = 0
    in_references = False
    ref_number_pattern = re.compile(r'^\[(\d+)\]\s*')
    references_found = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not in_references:
            if '参考文献' in text or '参 考 文 献' in text:
                if len(text) < 20:
                    in_references = True
                    print(f'\n找到参考文献标题 (段落 {i+1}): {text}')
                    continue
        else:
            if (text.startswith('致谢') or text.startswith('致 谢') or 
                text.startswith('结 论') or text.startswith('结论')):
                print(f'\n遇到结束标记，停止检测 (段落 {i+1}): {text}')
                break
            
            if not text:
                continue
            
            ref_match = ref_number_pattern.match(text)
            if ref_match:
                reference_count += 1
                ref_num = int(ref_match.group(1))
                references_found.append({
                    'number': ref_num,
                    'paragraph': i + 1,
                    'content': text
                })
                print(f'[{ref_num}] 段落 {i+1}: {text[:60]}...')
            else:
                has_ref_marker = False
                for marker in ['[J]', '[D]', '[M]', '[N]', '[P]', '[S]', '[EB/OL]']:
                    if marker in text:
                        has_ref_marker = True
                        break
                
                if has_ref_marker:
                    reference_count += 1
                    references_found.append({
                        'number': reference_count,
                        'paragraph': i + 1,
                        'content': text
                    })
                    print(f'[{reference_count}] 段落 {i+1}: {text[:60]}...')
    
    print(f'\n找到参考文献数量: {reference_count}')
    
    if references_found:
        print(f'\n已识别的参考文献编号: {[ref["number"] for ref in references_found]}')
        
        expected_num = 1
        missing_refs = []
        for ref in references_found:
            if ref['number'] != expected_num:
                while expected_num < ref['number']:
                    missing_refs.append(expected_num)
                    expected_num += 1
            expected_num += 1
        
        if missing_refs:
            print(f'⚠️  发现缺失的参考文献编号: {missing_refs}')
    
    if reference_count >= 6:
        print(f'✅ 参考文献数量检查通过！')
        return True
    else:
        print(f'❌ 参考文献数量不足，要求不低于6个，实际{reference_count}个')
        return False


def check_special_headings(doc):
    print('\n' + '=' * 80)
    print('【特殊标题格式检测】')
    print('=' * 80)
    
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    target_configs = [
        {'name': '结  论', 'patterns': ['结  论', '结 论', '结论']},
        {'name': '致 谢', 'patterns': ['致 谢', '致谢']},
        {'name': '参 考 文 献', 'patterns': ['参 考 文 献', '参 考 文 献', '参考文献']}
    ]
    
    all_paras = list(doc.paragraphs)
    
    for config in target_configs:
        found = False
        for i, para in enumerate(all_paras):
            text = para.text.strip()
            if any(text == pattern for pattern in config['patterns']):
                found = True
                print(f'\n检测到「{config["name"]}」(段落 {i+1}): {text}')
                
                para_elem = para._element
                pPr = para_elem.find('./w:pPr', namespaces=ns)
                
                jc_ok = False
                if pPr is not None:
                    jc = pPr.find('./w:jc', namespaces=ns)
                    if jc is not None:
                        val = jc.get(f'{{{ns["w"]}}}val')
                        if val == 'center':
                            jc_ok = True
                            print(f'  ✅ 对齐方式: 居中')
                        else:
                            errors.append({
                                'heading': config["name"],
                                'type': '对齐方式错误',
                                'actual': val,
                                'expected': 'center'
                            })
                            print(f'  ❌ 对齐方式: {val} (期望: center)')
                    else:
                        errors.append({
                            'heading': config["name"],
                            'type': '对齐方式错误',
                            'actual': '未设置',
                            'expected': 'center'
                        })
                        print(f'  ❌ 对齐方式: 未设置 (期望: center)')
                else:
                    errors.append({
                        'heading': config["name"],
                        'type': '对齐方式错误',
                        'actual': '未设置',
                        'expected': 'center'
                    })
                    print(f'  ❌ 对齐方式: 未设置 (期望: center)')
                
                font_ok = False
                sz_printed = False
                eastAsia_printed = False
                for run in para.runs:
                    if run.text.strip():
                        r_elem = run._element
                        rPr = r_elem.find('./w:rPr', namespaces=ns)
                        if rPr is not None:
                            sz = rPr.find('./w:sz', namespaces=ns)
                            rFonts = rPr.find('./w:rFonts', namespaces=ns)
                            
                            if sz is not None and not sz_printed:
                                sz_printed = True
                                sz_val = float(sz.get(f'{{{ns["w"]}}}val')) / 2.0
                                if abs(sz_val - 15.0) < 0.1:
                                    print(f'  ✅ 字体大小: {sz_val:.1f} 磅 (小3号)')
                                else:
                                    errors.append({
                                        'heading': config["name"],
                                        'type': '字体大小错误',
                                        'actual': f'{sz_val:.1f} 磅',
                                        'expected': '15.0 磅 (小3号)'
                                    })
                                    print(f'  ❌ 字体大小: {sz_val:.1f} 磅 (期望: 15.0 磅)')
                            
                            if rFonts is not None and not eastAsia_printed:
                                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                                if eastAsia:
                                    eastAsia_printed = True
                                    if eastAsia == '黑体':
                                        font_ok = True
                                        print(f'  ✅ 中文字体: 黑体')
                                    else:
                                        errors.append({
                                            'heading': config["name"],
                                            'type': '中文字体错误',
                                            'actual': eastAsia,
                                            'expected': '黑体'
                                        })
                                        print(f'  ❌ 中文字体: {eastAsia} (期望: 黑体)')
                                elif not eastAsia_printed:
                                    eastAsia_printed = True
                                    errors.append({
                                        'heading': config["name"],
                                        'type': '中文字体错误',
                                        'actual': '未设置',
                                        'expected': '黑体'
                                    })
                                    print(f'  ❌ 中文字体: 未设置 (期望: 黑体)')
                
                empty_lines_before = 0
                for j in range(i-1, max(0, i-10), -1):
                    if not all_paras[j].text.strip():
                        empty_lines_before += 1
                    else:
                        break
                
                if empty_lines_before >= 2:
                    print(f'  ✅ 段前空行数: {empty_lines_before} 行')
                else:
                    errors.append({
                        'heading': config["name"],
                        'type': '段前空行错误',
                        'actual': f'{empty_lines_before} 行',
                        'expected': '≥2 行'
                    })
                    print(f'  ❌ 段前空行数: {empty_lines_before} 行 (期望: ≥2 行)')
                
                break
        
        if not found:
            errors.append({
                'heading': config["name"],
                'type': '未找到标题',
                'actual': '未找到',
                'expected': '存在'
            })
            print(f'\n❌ 未找到「{config["name"]}」标题')
    
    if errors:
        print(f'\n发现 {len(errors)} 个特殊标题问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 「{error["heading"]}」:')
            print(f'   类型: {error["type"]}')
            print(f'   实际: {error["actual"]}')
            print(f'   期望: {error["expected"]}')
        return False
    else:
        print(f'\n✅ 所有特殊标题格式检查通过！')
        return True


def check_body_headings(doc):
    print('\n' + '=' * 80)
    print('【正文标题格式检测】')
    print('=' * 80)
    
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    h1_pattern = re.compile(r'^(\d+)\s+[\u4e00-\u9fa5]')
    h2_pattern = re.compile(r'^(\d+)\.(\d+)\s+[\u4e00-\u9fa5]')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        is_h1 = h1_pattern.match(text)
        is_h2 = h2_pattern.match(text)
        
        if not is_h1 and not is_h2:
            continue
        
        heading_level = '一级标题' if is_h1 else '二级标题'
        expected_font_size = 15.0 if is_h1 else 14.0
        print(f'\n检测到「{heading_level}」(段落 {i+1}): {text}')
        
        para_elem = para._element
        pPr = para_elem.find('./w:pPr', namespaces=ns)
        
        if is_h1:
            if pPr is not None:
                spacing = pPr.find('./w:spacing', namespaces=ns)
                if spacing is not None:
                    before = spacing.get(f'{{{ns["w"]}}}before')
                    after = spacing.get(f'{{{ns["w"]}}}after')
                    before_lines = float(before) / 240.0 if before else 0.0
                    after_lines = float(after) / 240.0 if after else 0.0
                    
                    if abs(before_lines - 0.5) < 0.1:
                        print(f'  ✅ 段前间距: {before_lines:.1f} 行')
                    else:
                        errors.append({
                            'paragraph': i + 1,
                            'text': text,
                            'heading_level': heading_level,
                            'type': '段前间距错误',
                            'actual': f'{before_lines:.1f} 行',
                            'expected': '0.5 行'
                        })
                        print(f'  ❌ 段前间距: {before_lines:.1f} 行 (期望: 0.5 行)')
                    
                    if abs(after_lines - 0.5) < 0.1:
                        print(f'  ✅ 段后间距: {after_lines:.1f} 行')
                    else:
                        errors.append({
                            'paragraph': i + 1,
                            'text': text,
                            'heading_level': heading_level,
                            'type': '段后间距错误',
                            'actual': f'{after_lines:.1f} 行',
                            'expected': '0.5 行'
                        })
                        print(f'  ❌ 段后间距: {after_lines:.1f} 行 (期望: 0.5 行)')
        
        font_ok = False
        sz_printed = False
        eastAsia_printed = False
        bold_ok = False
        bold_printed = False
        
        def is_style_bold(style):
            if style is None:
                return False
            style_elem = style._element
            rPr = style_elem.find('./w:rPr', namespaces=ns)
            if rPr is not None:
                b = rPr.find('./w:b', namespaces=ns)
                if b is not None:
                    b_val = b.get(f'{{{ns["w"]}}}val')
                    if b_val is None or b_val in ['1', 'true', 'on']:
                        return True
                bCs = rPr.find('./w:bCs', namespaces=ns)
                if bCs is not None:
                    bCs_val = bCs.get(f'{{{ns["w"]}}}val')
                    if bCs_val is None or bCs_val in ['1', 'true', 'on']:
                        return True
            if style.based_on:
                return is_style_bold(style.based_on)
            return False
        
        style_bold = False
        if para.style:
            style_bold = is_style_bold(para.style)
        
        for run in para.runs:
            if run.text.strip():
                r_elem = run._element
                rPr = r_elem.find('./w:rPr', namespaces=ns)
                if rPr is not None:
                    sz = rPr.find('./w:sz', namespaces=ns)
                    rFonts = rPr.find('./w:rFonts', namespaces=ns)
                    b = rPr.find('./w:b', namespaces=ns)
                    
                    if sz is not None and not sz_printed:
                        sz_printed = True
                        sz_val = float(sz.get(f'{{{ns["w"]}}}val')) / 2.0
                        if abs(sz_val - expected_font_size) < 0.1:
                            size_name = '小3号' if is_h1 else '4号'
                            print(f'  ✅ 字体大小: {sz_val:.1f} 磅 ({size_name})')
                        else:
                            size_name = '小3号' if is_h1 else '4号'
                            errors.append({
                                'paragraph': i + 1,
                                'text': text,
                                'heading_level': heading_level,
                                'type': '字体大小错误',
                                'actual': f'{sz_val:.1f} 磅',
                                'expected': f'{expected_font_size:.1f} 磅 ({size_name})'
                            })
                            print(f'  ❌ 字体大小: {sz_val:.1f} 磅 (期望: {expected_font_size:.1f} 磅 {size_name})')
                    
                    if rFonts is not None and not eastAsia_printed:
                        eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                        if eastAsia:
                            eastAsia_printed = True
                            if eastAsia == '黑体':
                                font_ok = True
                                print(f'  ✅ 中文字体: 黑体')
                            else:
                                errors.append({
                                    'paragraph': i + 1,
                                    'text': text,
                                    'heading_level': heading_level,
                                    'type': '中文字体错误',
                                    'actual': eastAsia,
                                    'expected': '黑体'
                                })
                                print(f'  ❌ 中文字体: {eastAsia} (期望: 黑体)')
                    
                    if not bold_printed:
                        bold_printed = True
                        is_bold = style_bold
                        if b is not None:
                            b_val = b.get(f'{{{ns["w"]}}}val')
                            if b_val is None or b_val in ['1', 'true', 'on']:
                                is_bold = True
                            elif b_val in ['0', 'false', 'off']:
                                is_bold = False
                        bCs = rPr.find('./w:bCs', namespaces=ns)
                        if bCs is not None:
                            bCs_val = bCs.get(f'{{{ns["w"]}}}val')
                            if bCs_val is None or bCs_val in ['1', 'true', 'on']:
                                is_bold = True
                        if is_bold:
                            bold_ok = True
                            print(f'  ✅ 加粗: 是')
                        else:
                            errors.append({
                                'paragraph': i + 1,
                                'text': text,
                                'heading_level': heading_level,
                                'type': '加粗错误',
                                'actual': '否',
                                'expected': '是'
                            })
                            print(f'  ❌ 加粗: 否 (期望: 是)')
    
    if errors:
        print(f'\n发现 {len(errors)} 个正文标题格式问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 段落 {error["paragraph"]} ({error["heading_level"]}):')
            print(f'   内容: {error["text"]}')
            print(f'   类型: {error["type"]}')
            print(f'   实际: {error["actual"]}')
            print(f'   期望: {error["expected"]}')
        return False
    else:
        print(f'\n✅ 所有正文标题格式检查通过！')
        return True


def check_reference_superscript(doc):
    print('\n' + '=' * 80)
    print('【参考文献引用上标检测】')
    print('=' * 80)
    
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    import re
    
    root = doc.element.body
    for i, p_elem in enumerate(root.findall('.//w:p', namespaces=ns)):
        text_parts = []
        for t in p_elem.findall('.//w:t', namespaces=ns):
            if t.text:
                text_parts.append(t.text)
        text = ''.join(text_parts)
        
        if re.search(r'\[\d+\]', text):
            if re.match(r'^\[\d+\]\s', text.strip()):
                continue
            
            r_elems = p_elem.findall('./w:r', namespaces=ns)
            for j, r_elem in enumerate(r_elems):
                t_elems = r_elem.findall('./w:t', namespaces=ns)
                run_text = ''.join([t.text for t in t_elems if t.text])
                
                if re.search(r'\[\d+\]', run_text):
                    rPr = r_elem.find('./w:rPr', namespaces=ns)
                    if rPr is not None:
                        vertAlign = rPr.find('./w:vertAlign', namespaces=ns)
                        if vertAlign is not None:
                            val = vertAlign.get(f'{{{ns["w"]}}}val')
                            if val != 'superscript':
                                errors.append({
                                    'paragraph': i + 1,
                                    'text': text[:100],
                                    'citation': run_text,
                                    'type': '上标格式错误',
                                    'actual': val or '无设置',
                                    'expected': 'superscript'
                                })
                        else:
                            errors.append({
                                'paragraph': i + 1,
                                'text': text[:100],
                                'citation': run_text,
                                'type': '缺少上标格式',
                                'actual': '无设置',
                                'expected': 'superscript'
                            })
    
    if errors:
        print(f'\n发现 {len(errors)} 个参考文献引用上标问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 段落 {error["paragraph"]}:')
            print(f'   内容: {error["text"]}...')
            print(f'   引用: {error["citation"]}')
            print(f'   类型: {error["type"]}')
            print(f'   实际: {error["actual"]}')
            print(f'   期望: {error["expected"]}')
    else:
        print(f'\n✅ 所有参考文献引用上标检查通过！')
    
    return len(errors) == 0


def check_reference_hyperlink(doc):
    print('\n' + '=' * 80)
    print('【参考文献引用交叉引用/超链接检测】')
    print('=' * 80)
    
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
    
    import re
    
    root = doc.element.body
    for i, p_elem in enumerate(root.findall('.//w:p', namespaces=ns)):
        text_parts = []
        for t in p_elem.findall('.//w:t', namespaces=ns):
            if t.text:
                text_parts.append(t.text)
        text = ''.join(text_parts)
        
        if re.search(r'\[\d+\]', text):
            if re.match(r'^\[\d+\]\s', text.strip()):
                continue
            
            found_valid_link = False
            
            # 检查超链接
            hyperlink_elems = p_elem.findall('.//w:hyperlink', namespaces=ns)
            if hyperlink_elems:
                for hyperlink_elem in hyperlink_elems:
                    text_elems = hyperlink_elem.findall('.//w:t', namespaces=ns)
                    hyperlink_text = ''.join([t.text for t in text_elems if t.text])
                    
                    if re.search(r'\[\d+\]', hyperlink_text):
                        rid = hyperlink_elem.get(f'{{{ns["r"]}}}id')
                        if rid:
                            found_valid_link = True
            
            # 检查交叉引用（域代码）
            if not found_valid_link:
                fld_simple = p_elem.findall('.//w:fldSimple', namespaces=ns)
                fld_char = p_elem.findall('.//w:fldChar', namespaces=ns)
                instr_text = p_elem.findall('.//w:instrText', namespaces=ns)
                
                if fld_simple or fld_char or instr_text:
                    for it in instr_text:
                        if it.text and ('REF' in it.text or 'PAGEREF' in it.text or 'NOTEREF' in it.text):
                            found_valid_link = True
                            break
            
            if not found_valid_link:
                errors.append({
                    'paragraph': i + 1,
                    'text': text[:100],
                    'type': '缺少交叉引用或超链接',
                    'message': '参考文献引用未设置交叉引用或超链接'
                })
    
    if errors:
        print(f'\n发现 {len(errors)} 个参考文献引用问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 段落 {error["paragraph"]}:')
            print(f'   内容: {error["text"]}...')
            print(f'   类型: {error["type"]}')
            print(f'   说明: {error["message"]}')
    else:
        print(f'\n✅ 所有参考文献引用交叉引用/超链接检查通过！')
    
    return len(errors) == 0


def check_keywords(doc):
    print('\n' + '=' * 80)
    print('【关键词格式检测】')
    print('=' * 80)
    
    errors = []
    found_keywords = False
    
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    doc_styles = {s.name: s for s in doc.styles}
    
    def get_eastasia_font_from_style(style, ns):
        if style is None:
            return None
        style_elem = style._element
        rPr = style_elem.find('./w:rPr', namespaces=ns)
        if rPr is not None:
            rFonts = rPr.find('./w:rFonts', namespaces=ns)
            if rFonts is not None:
                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                if eastAsia:
                    return eastAsia
        based_on = getattr(style, 'based_on', None)
        if based_on:
            return get_eastasia_font_from_style(based_on, ns)
        return None
    
    def get_run_font(run, para, doc_styles, ns):
        run_elem = run._element
        rPr = run_elem.find('./w:rPr', namespaces=ns)
        if rPr is not None:
            rFonts = rPr.find('./w:rFonts', namespaces=ns)
            if rFonts is not None:
                eastAsia = rFonts.get(f'{{{ns["w"]}}}eastAsia')
                if eastAsia:
                    return eastAsia
        if para.style:
            eastAsia_font = get_eastasia_font_from_style(para.style, ns)
            if eastAsia_font:
                return eastAsia_font
        default_style = doc_styles.get('Normal')
        if default_style:
            eastAsia_font = get_eastasia_font_from_style(default_style, ns)
            if eastAsia_font:
                return eastAsia_font
        return None
    
    def is_font_heiti(font_name):
        if not font_name:
            return False
        return ('黑体' in font_name or 
                'Heiti' in font_name or 
                'SimHei' in font_name or
                'Microsoft YaHei' in font_name or
                '微软雅黑' in font_name)
    
    def get_run_bold(run):
        if run.font.bold is not None:
            return run.font.bold
        return False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if '关键词' in text or 'Keywords' in text:
            found_keywords = True
            print(f'\n找到关键词段落 (段落 {i + 1}): {text}')
            
            for j, run in enumerate(para.runs):
                run_text = run.text
                if '关键词' in run_text:
                    print(f'\n  检测 \"关键词\" 文字:')
                    
                    current_font = get_run_font(run, para, doc_styles, ns)
                    print(f'    字体: {current_font or "未知"}')
                    
                    is_bold = get_run_bold(run)
                    print(f'    加粗: {is_bold}')
                    
                    if not is_font_heiti(current_font):
                        errors.append({
                            'type': '字体错误',
                            'paragraph': i + 1,
                            'text': f'"关键词"',
                            'actual': current_font or '未知',
                            'expected': '黑体'
                        })
                    
                    if is_bold:
                        errors.append({
                            'type': '加粗错误',
                            'paragraph': i + 1,
                            'text': f'"关键词"',
                            'actual': '加粗',
                            'expected': '不加粗'
                        })
            
            break
    
    if not found_keywords:
        print('\n⚠️  未找到关键词段落')
        errors.append({
            'type': '未找到',
            'message': '文档中未找到关键词段落'
        })
    
    if errors:
        print('\n发现的问题:')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. [{error["type"]}]')
            if 'paragraph' in error:
                print(f'   段落: {error["paragraph"]}')
            if 'text' in error:
                print(f'   内容: {error["text"]}')
            if 'actual' in error and 'expected' in error:
                print(f'   实际: {error["actual"]}')
                print(f'   期望: {error["expected"]}')
            if 'message' in error:
                print(f'   说明: {error["message"]}')
    else:
        print('\n✅ 关键词格式检查通过！')
    
    return len(errors) == 0


def check_empty_lines(doc):
    print('\n' + '=' * 80)
    print('【段后空行检测 - 详细信息】')
    print('=' * 80)
    
    errors = []
    all_paras = []
    check_para_indices = set()
    para_element_map = {}
    
    chinese_pattern = re.compile(r'[\u4e00-\u9fff]')
    
    body = doc._element.body
    elements = []
    for child in body:
        if child.tag.endswith('}p'):
            elements.append(('paragraph', child))
        elif child.tag.endswith('}tbl'):
            elements.append(('table', child))
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        has_page_break = False
        if para.paragraph_format.page_break_before:
            has_page_break = True
        else:
            for run in para.runs:
                if hasattr(run, '_element'):
                    elem = run._element
                    if elem.find('.//w:br', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                        has_page_break = True
                        break
        
        para_element_map[para._element] = i
        
        all_paras.append({
            'index': i,
            'text': text,
            'has_text': bool(text),
            'has_page_break': has_page_break,
            'element': para._element
        })
        
        is_check_para = True
        
        if i < 50:
            if (text.startswith('教育培训学院') or 
                text.startswith('2026届') or
                text.startswith('题    目') or
                text.startswith('学习形式') or
                text.startswith('学习层次') or
                text.startswith('专    业') or
                text.startswith('学生姓名') or
                text.startswith('学    号') or
                text.startswith('指导教师') or
                text.startswith('答辩日期') or
                text.startswith('评价数据的')):
                is_check_para = False
        
        if (text.startswith('第') and '章' in text) or \
           (text.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9')) and ' ' in text[:10]) or \
           (text.startswith('摘要') or text.startswith('Abstract') or 
            text.startswith('关键词') or text.startswith('Keywords') or
            text.startswith('参考文献') or text.startswith('参 考 文 献') or
            text.startswith('致 谢') or text.startswith('致谢') or
            text.startswith('目 次') or text.startswith('目次') or
            text.startswith('结 论') or text.startswith('结论') or
            text.startswith('摘  要') or text.startswith('目   次')):
            is_check_para = False
        
        if text.startswith('图 ') or text.startswith('表 ') or re.match(r'^\d+\.\d+', text):
            is_check_para = False
        
        if re.match(r'^\[\d+\]', text):
            is_check_para = False
        
        if not chinese_pattern.search(text) and not text.startswith('Testing'):
            is_check_para = False
        
        if is_check_para and text:
            check_para_indices.add(i)
    
    for i in range(len(all_paras)):
        current = all_paras[i]
        
        if current['index'] not in check_para_indices:
            continue
        
        if not current['has_text']:
            continue
        
        empty_count = 0
        found_next_text = False
        has_table_between = False
        
        current_element_idx = -1
        for idx, (elem_type, elem) in enumerate(elements):
            if elem_type == 'paragraph' and elem == current['element']:
                current_element_idx = idx
                break
        
        if current_element_idx == -1:
            continue
        
        j = current_element_idx + 1
        while j < len(elements):
            elem_type, elem = elements[j]
            
            if elem_type == 'table':
                has_table_between = True
                empty_count = 0
                break
            
            if elem_type == 'paragraph':
                para_idx = para_element_map.get(elem)
                if para_idx is not None:
                    para = all_paras[para_idx]
                    if para['has_text']:
                        found_next_text = True
                        break
                    else:
                        if not para['has_page_break']:
                            empty_count += 1
                        else:
                            print(f'  段落 {para_idx + 1} 是分页符，不计入空行')
            
            j += 1
        
        if found_next_text and not has_table_between and empty_count > 2:
            errors.append({
                'paragraph': current['index'] + 1,
                'text': current['text'][:80] + '...' if len(current['text']) > 80 else current['text'],
                'empty_lines': empty_count
            })
        elif has_table_between:
            print(f'  段落 {current["index"] + 1} 与下一段之间有表格，重新计算空行')
    
    print(f'\n检查了 {len(check_para_indices)} 个有效段落')
    print(f'发现 {len(errors)} 个段后空行问题')
    
    if errors:
        print(f'\n发现 {len(errors)} 个段后空行问题:\n')
        for i, error in enumerate(errors, 1):
            print(f'\n{i}. 段落 {error["paragraph"]}:')
            print(f'   内容: {error["text"]}')
            print(f'   段后空行数: {error["empty_lines"]} 行（超过2行）')
    else:
        print(f'\n✅ 所有段落段后空行检查通过！')
    
    return len(errors) == 0


def main():
    if len(sys.argv) > 1:
        docx_file = sys.argv[1]
        if not os.path.exists(docx_file):
            print(f'错误: 文件不存在 - {docx_file}')
            sys.exit(1)
    else:
        docx_files = glob.glob('*.docx')
        docx_files = [f for f in docx_files if not os.path.basename(f).startswith('~$')]
        if not docx_files:
            print('错误: 当前目录下未找到 .docx 文件')
            print('使用方法: python main.py [文档路径]')
            sys.exit(1)
        docx_file = docx_files[0]
    
    print(f'正在检查文档: {docx_file}')
    print('=' * 80)
    
    doc = Document(docx_file)
    
    fig_ok = check_figure_numbering(doc)
    ref_ok = check_references(doc)
    indent_ok = check_indent(doc)
    font_ok = check_chinese_font(doc)
    toc_heading1_ok = check_toc_heading1_font(doc)
    keywords_ok = check_keywords(doc)
    lines_ok = check_empty_lines(doc)
    superscript_ok = check_reference_superscript(doc)
    spacing_after_ok = check_paragraph_spacing_after(doc)
    body_font_ok = check_body_font_size(doc)
    table_font_ok = check_table_figure_font(doc)
    ref_count_ok = check_reference_count(doc)
    special_headings_ok = check_special_headings(doc)
    body_headings_ok = check_body_headings(doc)
    ref_crossref_ok = check_reference_crossref(doc)
    
    print('\n' + '=' * 80)
    print('【总体检测结果】')
    print('=' * 80)
    print(f'图序编号检测: {"✅ 通过" if fig_ok else "❌ 存在问题"}')
    print(f'参考文献格式: {"✅ 通过" if ref_ok else "❌ 存在问题"}')
    print(f'段落首行缩进: {"✅ 通过" if indent_ok else "❌ 存在问题"}')
    print(f'中文字体检测: {"✅ 通过" if font_ok else "❌ 存在问题"}')
    print(f'目录字体检测: {"✅ 通过" if toc_heading1_ok else "❌ 存在问题"}')
    print(f'关键词格式检测: {"✅ 通过" if keywords_ok else "❌ 存在问题"}')
    print(f'段后空行检测: {"✅ 通过" if lines_ok else "❌ 存在问题"}')
    print(f'参考文献上标检测: {"✅ 通过" if superscript_ok else "❌ 存在问题"}')
    print(f'段落段后间距: {"✅ 通过" if spacing_after_ok else "❌ 存在问题"}')
    print(f'正文字体行距检测: {"✅ 通过" if body_font_ok else "❌ 存在问题"}')
    print(f'表格插图字体检测: {"✅ 通过" if table_font_ok else "❌ 存在问题"}')
    print(f'参考文献数量检测: {"✅ 通过" if ref_count_ok else "❌ 存在问题"}')
    print(f'特殊标题格式检测: {"✅ 通过" if special_headings_ok else "❌ 存在问题"}')
    print(f'正文标题格式检测: {"✅ 通过" if body_headings_ok else "❌ 存在问题"}')
    print(f'参考文献交叉引用检测: {"✅ 通过" if ref_crossref_ok else "❌ 存在问题"}')
    print('\n' + '=' * 80)
    
    if fig_ok and ref_ok and indent_ok and font_ok and toc_heading1_ok and keywords_ok and lines_ok and superscript_ok and spacing_after_ok and body_font_ok and table_font_ok and ref_count_ok and special_headings_ok and body_headings_ok and ref_crossref_ok:
        print('🎉 所有检测项目均通过！')
    else:
        print('⚠️  部分检测项目存在问题，请查看上方详细信息')
    print('=' * 80)


if __name__ == '__main__':
    main()